import os
import sys
from docx import Document
import markdown
from bs4 import BeautifulSoup

current_dir = os.path.dirname(os.path.abspath(__file__))


def convert_md_to_docx(md_file_path: str, word_output_path: str) -> dict:
    # 从请求的内容中读取
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_str = f.read()
    html = markdown.markdown(md_str)
    soup = BeautifulSoup(html, 'html.parser')

    # Create a new Word document
    doc = Document(os.path.join(current_dir, "template.docx"))

    # 2. 清空所有段落（包括文本、表格、图片等）
    for element in list(doc.element.body):  # 遍历所有元素
        doc.element.body.remove(element)

    # Process each HTML element and add to the Word document
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol']):
        if element.name.startswith('h'):
            level = int(element.name[1])
            if level == 1:
                doc.add_paragraph(element.get_text().strip(), style="Title1")
            elif level == 2:
                doc.add_paragraph(element.get_text().strip(), style="Title2")
            else:
                doc.add_paragraph(element.get_text().strip(), style="Title3")
        elif element.name == 'p':
            paragraph_style = "Paragraph"
            paragraph_text = element.get_text().strip()
            # 每个换行符都添加一个段落
            for line in paragraph_text.split('\n'):
                doc.add_paragraph(line.strip(), style=paragraph_style)
        elif element.name == 'ol':
            for item in element.find_all('li'):
                doc.add_paragraph(item.get_text().strip(), style='NumberList')
        elif element.name == 'ul':
            for item in element.find_all('li'):
                doc.add_paragraph(item.get_text().strip(), style='BulletList')

    # 确保输出目录存在
    output_path = word_output_path
    target_dir = os.path.dirname(output_path)
    if target_dir and not os.path.exists(target_dir):
        os.makedirs(target_dir)

    doc.save(output_path)

    return {"docx_path": output_path}


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python md_to_docx.py <markdown_file_path> <word_output_path>")
        sys.exit(1)

    md_file_path = sys.argv[1]
    word_output_path = sys.argv[2]

    res = convert_md_to_docx(md_file_path, word_output_path)
    print(res["docx_path"])
